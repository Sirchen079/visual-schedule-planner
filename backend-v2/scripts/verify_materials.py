"""Frozen long-material reading and source-backed plans; temporary data only."""
# ruff: noqa: DTZ011 -- v2 uses local calendar dates.
import argparse
import http.client
import json
import sqlite3
import tempfile
from datetime import date, timedelta
from io import BytesIO
from pathlib import Path

from verify_ledger import request, start, stop


def upload(port, name, data):
    boundary = 'zhishi-material-check'
    body = (f'--{boundary}\r\nContent-Disposition: form-data; name="file"; filename="{name}"\r\n'
            'Content-Type: application/octet-stream\r\n\r\n').encode()+data+f'\r\n--{boundary}--\r\n'.encode()
    conn = http.client.HTTPConnection('127.0.0.1', port, timeout=30)
    try:
        conn.request('POST','/ai/attachments',body,{'Content-Type':f'multipart/form-data; boundary={boundary}'})
        response = conn.getresponse()
        raw = response.read()
        assert response.status == 201, raw[:500]
        return json.loads(raw)['file_id']
    finally:
        conn.close()


def pdf_bytes():
    objects, kids = [b'<< /Type /Catalog /Pages 2 0 R >>', b''], []
    for number in range(1, 9):
        page_id, content_id = len(objects)+1, len(objects)+2
        kids.append(f'{page_id} 0 R')
        objects.append(f'<< /Type /Page /Parent 2 0 R /MediaBox [0 0 600 800] /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> /Contents {content_id} 0 R >>'.encode())
        stream = f'BT /F1 12 Tf 50 750 Td (PAGE {number}: important conclusion) Tj ET'.encode()
        objects.append(f'<< /Length {len(stream)} >>\nstream\n'.encode()+stream+b'\nendstream')
    objects[1] = f'<< /Type /Pages /Kids [{" ".join(kids)}] /Count 8 >>'.encode()
    pdf, offsets = b'%PDF-1.4\n', [0]
    for number, obj in enumerate(objects, 1):
        offsets.append(len(pdf))
        pdf += f'{number} 0 obj\n'.encode()+obj+b'\nendobj\n'
    xref = len(pdf)
    pdf += f'xref\n0 {len(offsets)}\n0000000000 65535 f \n'.encode()
    pdf += b''.join(f'{offset:010d} 00000 n \n'.encode() for offset in offsets[1:])
    return pdf+f'trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF'.encode()


def verify(exe):
    root = Path(tempfile.mkdtemp(prefix='zhishi-material-frozen-'))
    print(f'MATERIAL_CHECK_ROOT={root}', flush=True)
    proc, log, port = start(exe, root, 'first')
    try:
        text = 'Background notes. '*3000+'Final requirement: reproduce the baseline before comparing methods.'
        fid = upload(port,'long.txt',text.encode())
        from urllib.parse import urlencode
        def search(fid, query):
            return request(port,'/api/materials/search?'+urlencode({'file_id':fid,'query':query}))
        hit = search(fid,'Final requirement')['hits'][0]
        doc = request(port,f"/api/materials/{fid}?part={hit['part']}&revision={hit['revision']}")
        quote = 'reproduce the baseline before comparing methods'
        cited_part = next(p['part'] for p in doc['parts'] if quote in p['text'])
        assert doc['document']['indexed_chars'] == len(text) and hit['part'] > 15
        request(port,f'/api/materials/{fid}?revision=stale',expected=409)
        pdf_id = upload(port,'eight.pdf',pdf_bytes())
        pdf = search(pdf_id,'PAGE 8')['hits']
        assert any(h['location']=='第 8 页' and 'PAGE 8' in h['excerpt'] for h in pdf)
        csv_id = upload(port,'rows.csv',('Item,Cost\n'+'Common,1\n'*100+'TailReceipt,42').encode())
        assert search(csv_id,'TailReceipt')['hits']
        from docx import Document
        word = Document()
        for i in range(100):
            word.add_paragraph(f'Background paragraph {i}. '*30)
        word.add_paragraph('WordTailConclusion')
        output = BytesIO(); word.save(output)
        assert search(upload(port,'long.docx',output.getvalue()),'WordTailConclusion')['hits']
        from openpyxl import Workbook
        workbook = Workbook()
        for i in range(22):
            workbook.create_sheet(f'Sheet{i}').append([f'LaterSheetContent{i}'])
        output = BytesIO(); workbook.save(output); workbook.close()
        assert search(upload(port,'sheets.xlsx',output.getvalue()),'LaterSheetContent21')['hits']
        p = request(port,'/api/research/projects','POST',{'title':'Material evidence project',
            'objective':'Reproduce the baseline','start_date':str(date.today()+timedelta(days=2))},201)
        pid = p['id']
        source = request(port,f'/api/research/projects/{pid}/materials','POST',{'file_id':fid},201)
        ref = {'source_id':source['id'],'part':cited_part,'revision':doc['document']['revision'],'quote':quote}
        payload = {'version':1,'rationale':'Follow the final requirement','steps':[{'title':'Reproduce',
            'outcome':'Save repeatable results','minutes':45,'source_refs':[ref]}]}
        bad = json.loads(json.dumps(payload)); bad['steps'][0]['source_refs'][0]['quote']='Invented quote'
        request(port,f'/api/research/projects/{pid}/plans','POST',bad,409)
        plan = request(port,f'/api/research/projects/{pid}/plans','POST',payload,201)
        request(port,f"/api/research/plans/{plan['id']}/apply",'POST')
        member = request(port,f'/api/research/projects/{pid}')['tasks'][0]
        assert member['source_refs'][0] == ref and quote in member['notes']
        revision = doc['document']['revision']
    finally:
        stop(proc,log,port)
    with sqlite3.connect(root/'v2/backend.db') as db:
        db.execute('UPDATE library_files SET extracted_text=?,content_sha256=NULL WHERE id=?',
            (json.dumps({'kind':'text','text':text[:30000],'tables':[]}),fid))
    proc, log, port = start(exe,root,'restart')
    try:
        doc = request(port,f'/api/materials/{fid}?part={cited_part}')
        assert doc['document']['revision'] == revision and any(quote in part['text'] for part in doc['parts'])
        detail = request(port,f'/api/research/projects/{pid}')
        assert detail['tasks'][0]['task_id'] == member['task_id']
        request(port,f'/api/files/{fid}','DELETE',expected=204)
        request(port,f'/api/materials/{fid}',expected=404)
        request(port,f'/api/files/{fid}/restore','POST')
        assert request(port,f'/api/materials/{fid}')['document']['revision'] == revision
        assert len(request(port,'/api/tasks')) == 1
        (root/'result.json').write_text(json.dumps({'ok':True,'file_id':fid,'project_id':pid,'revision':revision,
            'checks':['text tail','PDF page 8','CSV tail','Word tail','later workbook sheet','verified quote',
                      'forged quote rejected','restart and old cache upgrade','delete/restore']}),encoding='utf8')
    finally:
        stop(proc,log,port)
    print('MATERIAL_FROZEN_PASS: text/pdf/docx/csv/xlsx tail, source quote validation, actual task, restart, cache migration, delete and restore',flush=True)


if __name__ == '__main__':
    parser=argparse.ArgumentParser()
    parser.add_argument('exe',type=Path)
    verify(parser.parse_args().exe.resolve())
