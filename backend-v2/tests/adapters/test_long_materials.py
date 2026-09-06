import json

from zhishi.adapters import parsers


def test_text_full_blocks_survive_cache_and_limit_is_explicit(tmp_path, monkeypatch):
    path = tmp_path/'long.txt'
    path.write_text('前文。'*10000+'真正的最后一段', encoding='utf8')
    doc = parsers.ParsedDoc(**json.loads(parsers.parse_file(path).to_json()))
    assert '最后一段' in doc.blocks[-1]['text'] and len(doc.text) == parsers.MAX_CHARS
    monkeypatch.setattr(parsers, 'MAX_DOCUMENT_CHARS', 100)
    partial = parsers.parse_file(path)
    assert partial.partial and partial.warnings and sum(len(b['text']) for b in partial.blocks) == 100


def test_csv_and_xlsx_tail_rows_and_later_sheets_are_preserved(tmp_path):
    from openpyxl import Workbook
    csv = tmp_path/'rows.csv'
    csv.write_text('项目,金额\n'+'普通项目,1\n'*100+'尾部凭据,456.78', encoding='utf8')
    doc = parsers.parse_file(csv)
    assert len(doc.tables[0]) == 60 and '尾部凭据' in doc.blocks[-1]['text']
    assert '行 101' in doc.blocks[-1]['location']
    wb = Workbook()
    for i in range(22):
        sheet = wb.create_sheet('资料'+str(i))
        sheet.append(['第'+str(i)+'张表的内容'])
    path = tmp_path/'sheets.xlsx'
    wb.save(path)
    doc = parsers.parse_file(path)
    assert any('第21张表' in b['text'] and '资料21' in b['location'] for b in doc.blocks)


def test_docx_preserves_text_table_order_and_late_paragraphs(tmp_path):
    from docx import Document
    source = Document()
    source.add_paragraph('开头')
    table = source.add_table(rows=1, cols=1)
    table.cell(0,0).text = '中间的表格'
    for i in range(100):
        source.add_paragraph(f'第{i}段说明。' * 80)
    source.add_paragraph('正文最后的实际要求。')
    path = tmp_path/'long.docx'
    source.save(path)
    doc = parsers.parse_file(path)
    assert doc.blocks[0]['text'] == '开头' and doc.blocks[1]['text'] == '中间的表格'
    assert '实际要求' in doc.blocks[-1]['text'] and not doc.partial


def test_pdf_reads_after_fifth_page_and_reports_blank_scanned_pages(tmp_path):
    # Minimal real PDF, including its xref table; no optional PDF-creation dependency.
    objects = [b'<< /Type /Catalog /Pages 2 0 R >>', b'']
    kids = []
    for number in range(1, 9):
        page_id, content_id = len(objects)+1, len(objects)+2
        kids.append(f'{page_id} 0 R')
        objects.append(f'<< /Type /Page /Parent 2 0 R /MediaBox [0 0 600 800] /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> /Contents {content_id} 0 R >>'.encode())
        stream = f'BT /F1 12 Tf 50 750 Td (PAGE {number}: important conclusion) Tj ET'.encode() if number != 4 else b''
        objects.append(f'<< /Length {len(stream)} >>\nstream\n'.encode()+stream+b'\nendstream')
    objects[1] = f'<< /Type /Pages /Kids [{" ".join(kids)}] /Count 8 >>'.encode()
    pdf, offsets = b'%PDF-1.4\n', [0]
    for number, obj in enumerate(objects, 1):
        offsets.append(len(pdf))
        pdf += f'{number} 0 obj\n'.encode()+obj+b'\nendobj\n'
    xref = len(pdf)
    pdf += f'xref\n0 {len(offsets)}\n0000000000 65535 f \n'.encode()
    pdf += b''.join(f'{offset:010d} 00000 n \n'.encode() for offset in offsets[1:])
    pdf += f'trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF'.encode()
    path = tmp_path/'eight-pages.pdf'
    path.write_bytes(pdf)
    doc = parsers.parse_file(path)
    assert any('PAGE 8' in b['text'] and b['location'] == '第 8 页' for b in doc.blocks)
    assert doc.partial and any('第 4 页' in warning for warning in doc.warnings)
