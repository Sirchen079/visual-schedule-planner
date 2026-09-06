"""Structured previews plus bounded full-document blocks with original locations."""
from __future__ import annotations

import csv
import json
from dataclasses import asdict, dataclass, field
from itertools import islice
from pathlib import Path

PARSER_VERSION = 2
MAX_CHARS = 30_000
MAX_TABLES, MAX_ROWS = 20, 60
MAX_DOCUMENT_CHARS, MAX_PAGES, MAX_DOCUMENT_ROWS = 2_000_000, 500, 50_000
BLOCK_CHARS = 2000
IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp'}


@dataclass
class ParsedDoc:
    kind: str
    text: str = ''
    tables: list[list[list[str]]] = field(default_factory=list)
    blocks: list[dict] = field(default_factory=list)
    parser_version: int = PARSER_VERSION
    partial: bool = False
    warnings: list[str] = field(default_factory=list)

    def to_json(self) -> str:
        # Preview fields remain small; blocks hold the document body.
        if not self.blocks:
            builder = Blocks(self)
            builder.add('正文', self.text)
            for i, table in enumerate(self.tables):
                builder.add(f'表格 {i+1}', '\n'.join(' | '.join(row) for row in table))
        data = asdict(self)
        data['text'], data['tables'] = self.text[:MAX_CHARS], self.tables[:MAX_TABLES]
        return json.dumps(data, ensure_ascii=False)


class Blocks:
    def __init__(self, doc: ParsedDoc):
        self.doc, self.characters = doc, 0

    def warn(self, message: str):
        self.doc.partial = True
        if message not in self.doc.warnings:
            self.doc.warnings.append(message)

    def add(self, location: str, text: str):
        text = text.strip()
        available = MAX_DOCUMENT_CHARS - self.characters
        if len(text) > available:
            self.warn(f'解析正文达到 {MAX_DOCUMENT_CHARS} 字符上限，后续内容尚未处理。')
            text = text[:available]
        start = 0
        while start < len(text):
            end = min(start + BLOCK_CHARS, len(text))
            label = location if len(text) <= BLOCK_CHARS else f'{location} · 字符 {start+1}–{end}'
            self.doc.blocks.append({'location':label, 'text':text[start:end], 'overlap':120 if start else 0})
            if end == len(text):
                break
            start = end - 120
        self.characters += len(text)

    @property
    def full(self):
        return self.characters >= MAX_DOCUMENT_CHARS


def _trim(rows: list[list]) -> list[list[str]]:
    return [[('' if c is None else str(c)).strip() for c in row][:20] for row in rows[:MAX_ROWS]]


def _rows(builder: Blocks, rows, location: str) -> list[list[str]]:
    preview, batch = [], []
    first = 1
    for number, row in enumerate(rows, 1):
        if number > MAX_DOCUMENT_ROWS or builder.full:
            builder.warn(f'{location} 超出行数或正文容量上限，后续行尚未处理。')
            break
        values = [('' if value is None else str(value)).strip() for value in row]
        if len(preview) < MAX_ROWS:
            preview.append(values[:20])
        batch.append(' | '.join(values))
        if len(batch) == 20:
            builder.add(f'{location} · 行 {first}–{number}', '\n'.join(batch))
            batch, first = [], number+1
    if batch:
        builder.add(f'{location} · 行 {first}–{first+len(batch)-1}', '\n'.join(batch))
    return preview


def parse_file(path: Path) -> ParsedDoc:
    ext = path.suffix.lower()
    if ext in IMAGE_EXTS:
        return ParsedDoc(kind='image')
    if ext in ('.txt', '.md', '.log', '.json'):
        with path.open(encoding='utf-8-sig', errors='replace') as source:
            text = source.read(MAX_DOCUMENT_CHARS+1)
        doc = ParsedDoc(kind='text', text=text[:MAX_CHARS])
        Blocks(doc).add('正文', text)
        return doc
    if ext == '.csv':
        doc = ParsedDoc(kind='csv')
        with path.open(encoding='utf-8-sig', errors='replace', newline='') as source:
            doc.tables = [_rows(Blocks(doc), csv.reader(source), 'CSV')]
        return doc
    if ext == '.docx':
        return _parse_docx(path)
    if ext == '.xlsx':
        return _parse_xlsx(path)
    if ext == '.pdf':
        return _parse_pdf(path)
    return ParsedDoc(kind='unsupported')


def _parse_docx(path: Path) -> ParsedDoc:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    source, doc = Document(str(path)), ParsedDoc(kind='docx')
    builder, paragraphs, paragraph_number, table_number = Blocks(doc), [], 0, 0
    batch, first, preview_chars = [], 1, 0
    def flush():
        nonlocal batch, first
        if batch:
            builder.add(f'段落 {first}–{paragraph_number}', '\n'.join(batch))
            batch = []
        first = paragraph_number+1
    for element in source.element.body:
        if builder.full:
            builder.warn('正文容量已满，后续段落尚未处理。')
            break
        if element.tag.endswith('}p'):
            paragraph_number += 1
            text = Paragraph(element, source).text
            batch.append(text)
            if preview_chars < MAX_CHARS:
                paragraphs.append(text)
                preview_chars += len(text)
            if len(batch) >= 20:
                flush()
        elif element.tag.endswith('}tbl'):
            flush()
            table_number += 1
            table = Table(element, source)
            rows = _rows(builder, ([c.text for c in row.cells] for row in table.rows), f'表格 {table_number}')
            if len(doc.tables) < MAX_TABLES:
                doc.tables.append(rows)
    flush()
    doc.text = '\n'.join(paragraphs)[:MAX_CHARS]
    return doc


def _parse_xlsx(path: Path) -> ParsedDoc:
    from openpyxl import load_workbook
    source, doc = load_workbook(str(path), read_only=True, data_only=True), ParsedDoc(kind='xlsx')
    builder = Blocks(doc)
    try:
        for sheet in source.worksheets:
            if builder.full:
                builder.warn('正文容量已满，后续工作表尚未处理。')
                break
            preview = _rows(builder, sheet.iter_rows(values_only=True), f'工作表「{sheet.title}」')
            if len(doc.tables) < MAX_TABLES:
                doc.tables.append(preview)
    finally:
        source.close()
    return doc


def _parse_pdf(path: Path) -> ParsedDoc:
    import pdfplumber
    doc, texts, preview_chars = ParsedDoc(kind='pdf'), [], 0
    builder = Blocks(doc)
    with pdfplumber.open(str(path)) as source:
        if len(source.pages) > MAX_PAGES:
            builder.warn(f'PDF 超过 {MAX_PAGES} 页，后续页面尚未处理。')
        for number, page in enumerate(islice(source.pages, MAX_PAGES), 1):
            if builder.full:
                builder.warn('正文容量已满，后续页面尚未处理。')
                break
            text = page.extract_text() or ''
            builder.add(f'第 {number} 页', text)
            if not text.strip():
                builder.warn(f'第 {number} 页无可提取文本，可能需要视觉识别。')
            if preview_chars < MAX_CHARS:
                texts.append(text)
                preview_chars += len(text)
            if number <= 5:
                for table in (page.extract_tables() or [])[:MAX_TABLES-len(doc.tables)]:
                    doc.tables.append(_trim(table))
            page.close()
    doc.text = '\n'.join(texts)[:MAX_CHARS]
    return doc
